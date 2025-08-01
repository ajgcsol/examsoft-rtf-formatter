import tempfile
import os
import subprocess
import requests
from datetime import datetime
import streamlit as st
import pandas as pd
from pathlib import Path
import re
import docx
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import io
from io import BytesIO
import zipfile

# Try to import optional dependencies
try:
    from striprtf.striprtf import rtf_to_text
    RTF_AVAILABLE = True
except ImportError:
    RTF_AVAILABLE = False

# Import Microsoft 365 configuration
try:
    from examsoft_m365_config import M365_CONFIG
    SHAREPOINT_CONFIG_AVAILABLE = True
except ImportError:
    SHAREPOINT_CONFIG_AVAILABLE = False
    M365_CONFIG = {
        "client_id": "your-app-client-id",
        "tenant_id": "charlestonlaw.edu",
        "authority": "https://login.microsoftonline.com/charlestonlaw.edu",
        "scope": ["https://graph.microsoft.com/Sites.ReadWrite.All", 
                 "https://graph.microsoft.com/Files.ReadWrite.All"],
        "redirect_uri": "http://localhost:8501"
    }

# Import SharePoint integration module
try:
    from persistent_auth import initialize_persistent_auth, render_persistent_auth_ui, render_auth_status, sign_out_persistent
    # Define the upload function directly in this file to avoid import issues
    SHAREPOINT_INTEGRATION_AVAILABLE = True
    
    # Direct upload function with correct path
    def upload_to_sharepoint_corrected(access_token, file_content, filename):
        """Upload to correct SharePoint path: Exam Procedures/ExamSoft/Import with robust error handling"""
        try:
            import requests
            from urllib.parse import quote
            import re
            
            print("🔧 Using ROBUST upload function with enhanced debugging!")
            
            site_id = "charlestonlaw.sharepoint.com,ba0b6d09-2f32-4ccf-a24d-9a41e9be4a6a,ffe7f195-f2eb-4f68-af47-35a01fa9a2d7"
            
            # Step 1: Clean and validate filename
            # Remove invalid characters that might cause 400 errors
            clean_filename = re.sub(r'[<>:"/\\|?*\x00-\x1f]', '_', filename)
            clean_filename = re.sub(r'\s+', '_', clean_filename)  # Replace spaces with underscores
            clean_filename = clean_filename[:255]  # Limit length
            
            if clean_filename != filename:
                print(f"🔧 Cleaned filename: '{filename}' → '{clean_filename}'")
            
            # Step 2: Validate and prepare file content
            if isinstance(file_content, str):
                file_content = file_content.encode('utf-8')
                print(f"🔄 Converted string to bytes")
            
            if not isinstance(file_content, bytes):
                print(f"❌ Invalid content type: {type(file_content)}")
                return False, f"Invalid content type: {type(file_content)}"
            
            content_size = len(file_content)
            print(f"📊 Content size: {content_size:,} bytes")
            
            # Check for reasonable file size (SharePoint limit is 250GB, but let's be conservative)
            if content_size > 100 * 1024 * 1024:  # 100MB limit
                return False, f"File too large: {content_size:,} bytes (limit: 100MB)"
            
            if content_size == 0:
                return False, "Empty file content"
            
            # Step 3: Use different upload methods based on file size
            if content_size > 4 * 1024 * 1024:  # > 4MB, use session upload
                print("📤 Large file detected, using upload session...")
                return upload_large_file_with_session(site_id, access_token, file_content, clean_filename)
            else:
                print("📤 Small file, using direct upload...")
                return upload_small_file_direct(site_id, access_token, file_content, clean_filename)
                
        except Exception as e:
            print(f"❌ Exception in upload preparation: {str(e)}")
            import traceback
            traceback.print_exc()
            return False, f"Upload preparation error: {str(e)}"
    
    def upload_small_file_direct(site_id, access_token, file_content, filename):
        """Direct upload for small files using multiple API endpoint formats"""
        import requests
        from urllib.parse import quote
        
        encoded_filename = quote(filename)
        
        # Try different upload methods - the issue might be with the API endpoint format
        upload_methods = [
            {
                'name': 'Standard Path API',
                'url_template': 'https://graph.microsoft.com/v1.0/sites/{site_id}/drive/root:/{path}{filename}:/content',
                'paths': ['Exam Procedures/ExamSoft/Import/', 'ExamSoft/Import/', 'Import/', '']
            },
            {
                'name': 'Items API (no colons)',
                'url_template': 'https://graph.microsoft.com/v1.0/sites/{site_id}/drive/items/root:/{path}{filename}:/content',
                'paths': ['Exam Procedures/ExamSoft/Import/', 'ExamSoft/Import/', 'Import/', '']
            },
            {
                'name': 'Direct Drive API',
                'url_template': 'https://graph.microsoft.com/v1.0/drives/{drive_id}/root:/{path}{filename}:/content',
                'paths': ['Exam Procedures/ExamSoft/Import/', 'ExamSoft/Import/', 'Import/', ''],
                'needs_drive_id': True
            }
        ]
        
        # Get drive ID if needed
        drive_id = None
        
        for method in upload_methods:
            print(f"\n� Trying {method['name']}...")
            
            # Get drive ID for methods that need it
            if method.get('needs_drive_id') and not drive_id:
                try:
                    drive_response = requests.get(
                        f"https://graph.microsoft.com/v1.0/sites/{site_id}/drive",
                        headers={'Authorization': f'Bearer {access_token}', 'Accept': 'application/json'}
                    )
                    if drive_response.status_code == 200:
                        drive_id = drive_response.json().get('id')
                        print(f"   📋 Got drive ID: {drive_id}")
                    else:
                        print(f"   ❌ Cannot get drive ID: {drive_response.status_code}")
                        continue
                except Exception as e:
                    print(f"   ❌ Error getting drive ID: {e}")
                    continue
            
            for path in method['paths']:
                print(f"   � Path: '{path}'")
                
                # Build URL based on method
                if method.get('needs_drive_id'):
                    if not drive_id:
                        continue
                    upload_url = method['url_template'].format(
                        drive_id=drive_id, 
                        path=path, 
                        filename=encoded_filename
                    )
                else:
                    upload_url = method['url_template'].format(
                        site_id=site_id, 
                        path=path, 
                        filename=encoded_filename
                    )
                
                print(f"   � URL: {upload_url}")
                
                # Try with minimal headers first
                headers = {
                    'Authorization': f'Bearer {access_token}',
                    'Content-Type': 'application/octet-stream'
                }
                
                try:
                    response = requests.put(upload_url, headers=headers, data=file_content, timeout=60)
                    print(f"   📤 Response: {response.status_code}")
                    
                    if response.status_code in [200, 201]:
                        response_data = response.json()
                        file_url = response_data.get('webUrl', 'URL not available')
                        
                        print(f"✅ SUCCESS with {method['name']}!")
                        print(f"✅ Path: {path}")
                        print(f"🔗 File URL: {file_url}")
                        
                        return True, {
                            'url': file_url,
                            'size': response_data.get('size', 0),
                            'created': response_data.get('createdDateTime', 'N/A'),
                            'message': f'Successfully uploaded using {method["name"]} to {path or "root"}'
                        }
                    elif response.status_code == 400:
                        try:
                            error_data = response.json()
                            error_details = error_data.get('error', {})
                            error_code = error_details.get('code', 'Unknown')
                            error_message = error_details.get('message', 'Unknown error')
                            print(f"   ❌ 400 Error - Code: {error_code}, Message: {error_message}")
                            
                            # Check for specific errors that might give us clues
                            if 'malformed' in error_message.lower() or 'syntax' in error_message.lower():
                                print(f"   💡 Hint: URL syntax issue detected")
                            elif 'token' in error_message.lower():
                                print(f"   💡 Hint: Authentication issue detected")
                                return False, "Authentication token issue - please re-authenticate"
                        except:
                            print(f"   ❌ 400 Error - Raw: {response.text[:200]}")
                    elif response.status_code == 404:
                        print(f"   📂 Path not found - trying next...")
                    elif response.status_code == 403:
                        print(f"   🔒 Access denied - permissions issue")
                    else:
                        print(f"   ❌ Error {response.status_code}: {response.text[:100]}")
                        
                except requests.exceptions.Timeout:
                    print(f"   ⏱️ Timeout")
                except Exception as e:
                    print(f"   ❌ Exception: {str(e)}")
        
        # If everything fails, try a completely different approach - create a simple test file
        print(f"\n🧪 All methods failed. Trying simple test upload to root...")
        return try_simple_test_upload(site_id, access_token)
    
    def try_simple_test_upload(site_id, access_token):
        """Try uploading a simple test file to see if the API works at all"""
        import requests
        
        test_content = b"Test file content from ExamSoft formatter"
        test_filename = "test_upload.txt"
        
        # Try the simplest possible upload
        simple_urls = [
            f"https://graph.microsoft.com/v1.0/sites/{site_id}/drive/root/children/{test_filename}/content",
            f"https://graph.microsoft.com/v1.0/sites/{site_id}/drive/root:/{test_filename}:/content"
        ]
        
        headers = {
            'Authorization': f'Bearer {access_token}',
            'Content-Type': 'text/plain'
        }
        
        for url in simple_urls:
            print(f"   🧪 Testing: {url}")
            try:
                response = requests.put(url, headers=headers, data=test_content)
                print(f"   📤 Response: {response.status_code}")
                
                if response.status_code in [200, 201]:
                    print(f"   ✅ Simple upload works! The issue is with your RTF file or filename.")
                    return False, "API works but there's an issue with your specific file. Check filename characters and content format."
                else:
                    try:
                        error_data = response.json()
                        print(f"   ❌ Error: {error_data}")
                    except:
                        print(f"   ❌ Raw error: {response.text[:200]}")
            except Exception as e:
                print(f"   ❌ Exception: {e}")
        
        return False, "All upload methods failed - this may be a permissions or configuration issue"
    
    def upload_large_file_with_session(site_id, access_token, file_content, filename):
        """Upload large files using upload session"""
        import requests
        from urllib.parse import quote
        
        print(f"📤 Initiating upload session for large file...")
        
        encoded_filename = quote(filename)
        folder_path = "Exam Procedures/ExamSoft/Import"
        
        # Create upload session
        session_url = f"https://graph.microsoft.com/v1.0/sites/{site_id}/drive/root:/{folder_path}/{encoded_filename}:/createUploadSession"
        
        headers = {
            'Authorization': f'Bearer {access_token}',
            'Content-Type': 'application/json'
        }
        
        session_data = {
            "item": {
                "@microsoft.graph.conflictBehavior": "replace"
            }
        }
        
        try:
            session_response = requests.post(session_url, headers=headers, json=session_data)
            print(f"� Session creation: {session_response.status_code}")
            
            if session_response.status_code in [200, 201]:
                session_info = session_response.json()
                upload_url = session_info['uploadUrl']
                
                # Upload in chunks
                chunk_size = 327680  # 320KB chunks (must be multiple of 327680)
                total_size = len(file_content)
                
                for start in range(0, total_size, chunk_size):
                    end = min(start + chunk_size, total_size)
                    chunk = file_content[start:end]
                    
                    chunk_headers = {
                        'Content-Length': str(len(chunk)),
                        'Content-Range': f'bytes {start}-{end-1}/{total_size}'
                    }
                    
                    chunk_response = requests.put(upload_url, headers=chunk_headers, data=chunk)
                    print(f"📦 Chunk {start//chunk_size + 1}: {chunk_response.status_code}")
                    
                    if chunk_response.status_code not in [202, 200, 201]:
                        print(f"❌ Chunk upload failed: {chunk_response.text}")
                        return False, f"Chunk upload failed: {chunk_response.status_code}"
                
                # Final response should contain file info
                if chunk_response.status_code in [200, 201]:
                    response_data = chunk_response.json()
                    return True, {
                        'url': response_data.get('webUrl', 'URL not available'),
                        'size': response_data.get('size', 0),
                        'created': response_data.get('createdDateTime', 'N/A'),
                        'message': 'Successfully uploaded large file'
                    }
                else:
                    return False, f"Upload session completed but final status unclear: {chunk_response.status_code}"
            else:
                print(f"❌ Session creation failed: {session_response.text}")
                return False, f"Could not create upload session: {session_response.status_code}"
                
        except Exception as e:
            print(f"❌ Upload session error: {str(e)}")
            return False, f"Upload session error: {str(e)}"
        
        except Exception as upload_error:
            print(f"❌ SharePoint upload function error: {str(upload_error)}")
            return False, f"Upload function error: {str(upload_error)}"
            
except ImportError as e:
    SHAREPOINT_INTEGRATION_AVAILABLE = False

def generate_docx_with_questions(questions_list, instructions_text, output_path):
    """Generate a DOCX file with instructions and questions, formatted simply."""
    doc = Document()
    # Set default font to Times New Roman 12pt
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = docx.shared.Pt(12)
    # Do NOT add instructions to the main exam file
    # Add questions
    for q in questions_list:
        lines = q.split('\n')
        is_essay = False
        for idx, line in enumerate(lines):
            line = line.strip()
            # Replace problematic bullets/question marks with standard text
            line = line.replace('\u2022', '-')
            line = line.replace('\u25CF', '-')
            line = line.replace('\u25A0', '-')
            line = line.replace('\u25CB', '-')
            line = line.replace('\u25AA', '-')
            line = line.replace('\u25B2', '-')
            line = line.replace('\u25BA', '-')
            line = line.replace('\u25C6', '-')
            line = line.replace('\u25CF', '-')
            line = line.replace('\uFFFD', '-')
            # Remove any leading non-ASCII chars
            line = re.sub(r'^[^\x00-\x7F]+', '', line)
            if line == 'Type: E':
                is_essay = True
                # Do not add a separate heading, handled below
            elif is_essay and re.match(r'^\d+\. ', line):
                # Essay question number and content: prepend 'Type: E'
                qnum_match = re.match(r'^(\d+)(\. )(.*)', line)
                if qnum_match:
                    para = doc.add_paragraph()
                    run_type = para.add_run('Type: E ')
                    run_type.bold = True
                    run_type.font.name = 'Times New Roman'
                    run_type.font.size = docx.shared.Pt(12)
                    run_num = para.add_run(qnum_match.group(1) + qnum_match.group(2))
                    run_num.bold = True
                    run_num.font.name = 'Times New Roman'
                    run_num.font.size = docx.shared.Pt(12)
                    run_txt = para.add_run(qnum_match.group(3))
                    run_txt.font.name = 'Times New Roman'
                    run_txt.font.size = docx.shared.Pt(12)
                else:
                    para = doc.add_paragraph('Type: E ' + line)
                    for run in para.runs:
                        run.font.name = 'Times New Roman'
                        run.font.size = docx.shared.Pt(12)
            elif re.match(r'^\d+\. ', line):
                # Bold question number, normal text after (non-essay)
                qnum_match = re.match(r'^(\d+)(\. )(.*)', line)
                if qnum_match:
                    para = doc.add_paragraph()
                    run_num = para.add_run(qnum_match.group(1) + qnum_match.group(2))
                    run_num.bold = True
                    run_num.font.name = 'Times New Roman'
                    run_num.font.size = docx.shared.Pt(12)
                    run_txt = para.add_run(qnum_match.group(3))
                    run_txt.font.name = 'Times New Roman'
                    run_txt.font.size = docx.shared.Pt(12)
                else:
                    doc.add_paragraph(line)
            elif re.match(r'^[*]?[A-D]\. ', line):
                # Indented plain paragraph for answer choices, preserve asterisk
                para = doc.add_paragraph(line)
                para.paragraph_format.left_indent = docx.shared.Inches(0.25)
                for run in para.runs:
                    run.font.name = 'Times New Roman'
                    run.font.size = docx.shared.Pt(12)
            elif is_essay and idx > 0:
                # Essay content
                para = doc.add_paragraph(line)
                para.paragraph_format.left_indent = docx.shared.Inches(0.25)
                for run in para.runs:
                    run.font.name = 'Times New Roman'
                    run.font.size = docx.shared.Pt(12)
            else:
                para = doc.add_paragraph(line)
                for run in para.runs:
                    run.font.name = 'Times New Roman'
                    run.font.size = docx.shared.Pt(12)
        doc.add_paragraph('')
    doc.save(output_path)

def convert_docx_to_rtf_with_libreoffice(docx_path, rtf_path):
    """Convert DOCX to RTF using LibreOffice headless mode."""
    # LibreOffice outputs to the same directory as docx_path
    output_dir = os.path.dirname(rtf_path)
    cmd = [
        r'C:\\Program Files\\LibreOffice\\program\\soffice.exe', '--headless', '--convert-to', 'rtf', docx_path, '--outdir', output_dir
    ]
    subprocess.run(cmd, check=True)
    # LibreOffice names the output file with .rtf extension
    # Move/rename if needed
    generated_rtf = os.path.join(output_dir, os.path.splitext(os.path.basename(docx_path))[0] + '.rtf')
    if generated_rtf != rtf_path:
        os.replace(generated_rtf, rtf_path)

try:
    from sharepoint_integration_fixed import render_sharepoint_ui
    SHAREPOINT_AVAILABLE = True
except ImportError:
    SHAREPOINT_AVAILABLE = False

# Import Azure configuration loader
try:
    from azure_config_loader import get_converter_endpoint, show_azure_status, is_using_azure
    AZURE_CONFIG_AVAILABLE = True
except ImportError:
    AZURE_CONFIG_AVAILABLE = False
    # Fallback function
    def get_converter_endpoint():
        return "http://localhost:8080/convert"


def convert_docx_to_rtf_via_api(docx_path, rtf_path, api_url=None):
    """Send DOCX to the LibreOffice Docker API and save the returned RTF."""
    if api_url is None:
        api_url = get_converter_endpoint()
    
    with open(docx_path, "rb") as f:
        files = {'file': f}
        response = requests.post(api_url, files=files)
        response.raise_for_status()
        with open(rtf_path, "wb") as out:
            out.write(response.content)

def generate_filename(course, section, professor, file_type, extension="rtf"):
    """Generate filename with format: COURSE_SECTION_PROFESSOR_TYPE_YYMMDD.ext"""
    # Get current date in YYMMDD format
    date_stamp = datetime.now().strftime("%y%m%d")
    
    # Clean inputs
    course_clean = course.strip().upper() if course.strip() else "COURSE"
    section_clean = section.strip() if section.strip() else "001"
    professor_clean = professor.strip().title() if professor.strip() else "Professor"
    
    # Generate filename
    filename = f"{course_clean}_{section_clean}_{professor_clean}_{file_type}_{date_stamp}.{extension}"
    return filename

# SharePoint integration imports
try:
    import msal
    import json
    SHAREPOINT_AVAILABLE = True
except ImportError:
    SHAREPOINT_AVAILABLE = False

def clean_text_encoding(text):
    """Fix common mojibake/encoding issues from Windows-1252/UTF-8 mismatches."""
    replacements = {
        'â€™': "'",
        'â€œ': '"',
        'â€': '"',
        'â€˜': "'",
        'â€“': '-',
        'â€”': '—',
        'â€¦': '...',
        'â€¢': '•',
        'â€': '"',
        'â€\x9d': '"',
        'â€\x9c': '"',
        'â€\x98': "'",
        'â€\x99': "'",
        'Ã©': 'é',
        'Ã¨': 'è',
        'Ã¢': 'â',
        'Ãª': 'ê',
        'Ã®': 'î',
        'Ã´': 'ô',
        'Ã¶': 'ö',
        'Ã¼': 'ü',
        'Ã«': 'ë',
        'Ã§': 'ç',
        'Ã ': 'à',
        'Ã¹': 'ù',
        'Ã»': 'û',
        'Ã¼': 'ü',
        'ÃŸ': 'ß',
        'Â©': '©',
        'Â®': '®',
        'Â±': '±',
        'Â·': '·',
        'Â°': '°',
        'Â¼': '¼',
        'Â½': '½',
        'Â¾': '¾',
        'Â«': '«',
        'Â»': '»',
        'Â·': '·',
        'Â': '',
    }
    for bad, good in replacements.items():
        text = text.replace(bad, good)
    return text

def extract_text_from_rtf(file):
    """Extract clean text from RTF file"""
    try:
        from striprtf.striprtf import rtf_to_text
        content = file.read().decode('utf-8', errors='ignore')
        text = rtf_to_text(content)
        return clean_text_encoding(text)
    except ImportError:
        # Manual RTF text extraction
        content = file.read().decode('utf-8', errors='ignore')
        # ...existing code...
        # Remove any remaining hex/numeric artifacts
        content = re.sub(r'\b\d{8,}\b', ' ', content)
        text = content.strip()
        return clean_text_encoding(text)

def extract_text_from_docx(file):
    """Extract text from DOCX file"""
    doc = docx.Document(file)
    text = "\n".join([para.text for para in doc.paragraphs])
    return clean_text_encoding(text)

def extract_text(file):
    """Extract text from uploaded file based on file type"""
    file_extension = file.name.lower().split('.')[-1]
    if file_extension == 'rtf':
        return extract_text_from_rtf(file)
    elif file_extension in ['docx', 'doc']:
        return extract_text_from_docx(file)
    else:
        st.error(f"Unsupported file type: {file_extension}")
        return ""

def load_answer_key(file):
    """Load answer key from Excel or CSV file"""
    try:
        if file.name.endswith('.xlsx'):
            # Try reading with and without header
            try:
                df = pd.read_excel(file, header=0)
            except Exception:
                file.seek(0)
                df = pd.read_excel(file, header=None)
        else:
            try:
                df = pd.read_csv(file, header=0)
            except Exception:
                file.seek(0)
                df = pd.read_csv(file, header=None)

        # Find the first non-empty column (works for header/no-header)
        for col in df.columns:
            col_data = df[col].dropna().astype(str).str.strip()
            # Only use if at least half the rows are non-empty and look like answers (A-D or a-d)
            if col_data.str.match(r'^[A-Da-d]$').sum() >= max(1, len(col_data) // 2):
                answers = col_data.tolist()
                return answers
        # Fallback: just use the first column, skipping empty rows
        answers = df.iloc[:, 0].dropna().astype(str).str.strip().tolist()
        return answers
    except Exception as e:
        st.error(f"Error loading answer key: {e}")
        return []

def parse_exam_content(text):
    """Parse exam text into instructions and questions sections"""
    # text is already cleaned by clean_text_encoding
    # Find the split between instructions and questions
    multiple_choice_match = re.search(r'MULTIPLE[-\s]*CHOICE', text, re.IGNORECASE)
    if multiple_choice_match:
        instructions_text = text[:multiple_choice_match.start()].strip()
        questions_text = text[multiple_choice_match.end():].strip()
        return instructions_text, questions_text
    else:
        # If no MULTIPLE-CHOICE found, assume everything is questions
        return "", text

def format_instructions_rtf(text):
    """Format instructions as clean RTF"""
    # Clean the text
    text = re.sub(r'\s+', ' ', text)
    text = text.strip()
    
    # Create simple RTF format
    rtf_header = r"{\rtf1\ansi\deff0{\fonttbl{\f0 Times New Roman;}}\f0\fs24"
    
    # Convert to RTF with proper line breaks
    rtf_text = text.replace('\n', r'\line ')
    
    # Close RTF
    return rtf_header + rtf_text + "}"

def generate_instructions_docx(instructions_text):
    """Generate a DOCX file with instructions, return as bytes"""
    doc = Document()
    
    # Set default font to Times New Roman 12pt
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = docx.shared.Pt(12)
    
    # Add title
    title = doc.add_heading('INSTRUCTIONS', level=1)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    title_run = title.runs[0]
    title_run.font.name = 'Times New Roman'
    title_run.font.size = docx.shared.Pt(12)
    title_run.bold = True
    
    # Headers that should be bold
    bold_headers = ["Time", "Reference Materials", "Questions", "Multiple-Choice:", "Essay:", "Word Count:", "Please be sure to follow"]
    
    # Split instructions into paragraphs and format properly
    paragraphs = instructions_text.split('\n\n')
    for para_text in paragraphs:
        if para_text.strip():
            # Clean the text and handle encoding issues
            clean_para = clean_text_encoding(para_text.strip())
            para = doc.add_paragraph()
            
            # Check if this paragraph starts with a bold header
            is_bold_header = any(clean_para.startswith(header) for header in bold_headers)
            
            if is_bold_header:
                # Find where the header ends
                for header in bold_headers:
                    if clean_para.startswith(header):
                        # Add the header as bold
                        header_run = para.add_run(header)
                        header_run.font.name = 'Times New Roman'
                        header_run.font.size = docx.shared.Pt(12)
                        header_run.bold = True
                        
                        # Add the rest as normal text
                        remaining_text = clean_para[len(header):]
                        if remaining_text:
                            normal_run = para.add_run(remaining_text)
                            normal_run.font.name = 'Times New Roman'
                            normal_run.font.size = docx.shared.Pt(12)
                        break
            else:
                # Regular paragraph
                run = para.add_run(clean_para)
                run.font.name = 'Times New Roman'
                run.font.size = docx.shared.Pt(12)
    
    # Save to bytes
    doc_bytes = io.BytesIO()
    doc.save(doc_bytes)
    doc_bytes.seek(0)
    return doc_bytes.getvalue()

def parse_questions_from_text(questions_text, answer_key, use_asterisk_method=True):
    """Parse questions from the text and format for ExamSoft"""
    formatted_questions = []
    # Split on question numbers at start of line or start of text
    question_blocks = re.split(r'(?:^|\n)\s*(?:Type:\s*E\s+)?(\d+)\.\s+', questions_text)
    mc_index = 0  # index for multiple choice questions only
    for i in range(1, len(question_blocks), 2):
        if i + 1 < len(question_blocks):
            q_num = int(question_blocks[i])
            q_content = question_blocks[i + 1].strip()
            q_type = classify_question(q_content)
            if q_type == 'essay':
                formatted_q = format_essay_question(q_num, q_content)
            else:
                formatted_q = format_multiple_choice_question(q_num, q_content, answer_key, use_asterisk_method, mc_index)
                mc_index += 1
            if formatted_q:
                formatted_questions.append(formatted_q)
    return formatted_questions

def classify_question(q_content):
    """Classify question as 'mc' or 'essay' based on content."""
    # If any line matches A.–D. pattern, it's MC
    lines = q_content.split('\n')
    for line in lines:
        if re.match(r'^[A-D]\.\s+', line.strip()):
            return 'mc'
    # If 'ESSAY' or 'Type: E' present, or no answer choices, treat as essay
    if 'ESSAY' in q_content.upper() or re.search(r'Type:\s*E', q_content, re.IGNORECASE):
        return 'essay'
    # If no answer choices, treat as essay
    return 'essay'

def format_multiple_choice_question(q_num, content, answer_key, use_asterisk_method=True, mc_index=None):
    """Format a multiple choice question according to ExamSoft RTF import guidelines"""
    lines = content.split('\n')
    question_text = []
    choices = []
    current_section = "question"
    for line in lines:
        line = line.strip()
        if not line:
            continue
        choice_match = re.match(r'^([A-D])\.\s*(.+)', line)
        if choice_match:
            current_section = "choices"
            choice_letter = choice_match.group(1)
            choice_text = choice_match.group(2)
            # Check if this is the correct answer
            if use_asterisk_method and mc_index is not None and mc_index < len(answer_key):
                correct_answer = answer_key[mc_index].upper().strip()
                if choice_letter.upper() == correct_answer:
                    choices.append(f"*{choice_letter}. {choice_text}")
                else:
                    choices.append(f"{choice_letter}. {choice_text}")
            else:
                choices.append(f"{choice_letter}. {choice_text}")
        else:
            if current_section == "question":
                question_text.append(line)
    result = []
    if question_text:
        full_question = f"{q_num}. {' '.join(question_text)}"
        if '\n\n' in ' '.join(question_text):
            full_question = full_question.replace('\n\n', '<br>')
        result.append(full_question)
    result.extend(choices)
    return '\n'.join(result)

def format_essay_question(q_num, content):
    """Format an essay question according to ExamSoft RTF import guidelines"""
    # Clean up the content
    content = re.sub(r'Type:\s*E\s*\d*\.?\s*', '', content)
    content = content.strip()
    result = []
    # Split into paragraphs by double newlines
    paragraphs = [p.strip() for p in content.split('\n\n') if p.strip()]
    # Output 'Type: E' and question number/title outside <p> tags
    result.append(f"Type: E")
    # If first paragraph looks like a title (e.g. 'ESSAY' or has word count), put it as the question line
    if paragraphs:
        # If first paragraph is all-caps or contains 'ESSAY' or 'Word Count', treat as title
        first_para = paragraphs[0]
        if re.match(r'^[A-Z\s()\-:;0-9.]+$', first_para) or 'ESSAY' in first_para.upper() or 'WORD COUNT' in first_para.upper():
            result.append(f"{q_num}. {first_para}")
            para_body = paragraphs[1:]
        else:
            result.append(f"{q_num}.")
            para_body = paragraphs
        # Wrap each body paragraph in <p> tags
        for p in para_body:
            result.append(f"<p>{p}</p>")
    else:
        # Fallback: just output question number and single paragraph
        result.append(f"{q_num}.")
        result.append(f"<p>{content}</p>")
    return '\n'.join(result)

def create_rtf_content(questions_list, answer_key=None, use_answer_key_method=False):
    """Create RTF content for ExamSoft import"""
    # Use Times New Roman, 12pt, normal spacing, and proper paragraph breaks
    rtf_header = r"{\rtf1\ansi\deff0{\fonttbl{\f0 Times New Roman;}}\f0\fs24 "
    rtf_body = ""
    for question in questions_list:
        # Convert HTML <br> tags to RTF paragraph breaks
        question_rtf = question.replace("<br>", r"\par ")
        # Remove any other HTML tags
        question_rtf = re.sub(r'<[^>]+>', '', question_rtf)
        # Each question and its answers are separated by a single paragraph break
        rtf_body += question_rtf + r"\par "
    # Add answer key section if using the alternative method
    if use_answer_key_method and answer_key:
        rtf_body += r"\par Answers:\par "
        for i, answer in enumerate(answer_key, 1):
            rtf_body += f"{i}. {answer.lower()}\\par "
    return rtf_header + rtf_body + "}"


# Main Streamlit UI
st.subheader("Charleston School of Law")
st.header("ExamSoft RTF Formatter")
st.write("Paste your instructions, exam questions, and answer key below. No file upload needed.")

# Initialize SharePoint authentication early for better UX
if SHAREPOINT_INTEGRATION_AVAILABLE:
    try:
        initialize_persistent_auth()
        # Show authentication status in sidebar
        with st.sidebar:
            st.header("🔐 Microsoft 365")
            
            if render_persistent_auth_ui():
                # User is authenticated
                render_auth_status()
                
                # Add sign out button in sidebar too
                if st.button("🔓 Sign Out", key="sidebar_signout"):
                    if sign_out_persistent():
                        st.success("👋 Signed out!")
                        # Don't use st.rerun() - let natural refresh handle it
            else:
                st.info("🔑 Sign in for SharePoint upload")
                
    except Exception as e:
        st.sidebar.error(f"Auth error: {e}")

# Show Azure/Docker endpoint status
if AZURE_CONFIG_AVAILABLE:
    try:
        show_azure_status()
    except Exception as e:
        st.error(f"Azure config error: {e}")
        st.info("🏠 **Using Local Docker Endpoint** - Deploy to Azure for cloud conversion service")
else:
    st.info("🏠 **Using Local Docker Endpoint** - Deploy to Azure for cloud conversion service")# File naming inputs
col1, col2, col3 = st.columns(3)
with col1:
    course_input = st.text_input("Course", placeholder="e.g., CONST", help="Course abbreviation for filename")
with col2:
    section_input = st.text_input("Section", placeholder="e.g., 001", help="Section number for filename")
with col3:
    professor_input = st.text_input("Professor Last Name", placeholder="e.g., Smith", help="Professor's last name for filename")

st.subheader("Instructions (Optional)")
instructions_input = st.text_area("Paste Instructions", height=150, help="Optional: Paste or type the exam instructions. Leave blank if not needed.", placeholder="Instructions are optional - you can leave this blank if you only need to format questions.")

st.subheader("Exam Questions (Required - Paste all questions, including numbering and choices)")
questions_input = st.text_area("Paste Exam Questions", height=400, help="Paste all exam questions as plain text, including numbering and answer choices.")

st.subheader("Answer Key (Paste column from Excel, one answer per line, e.g. A B C D ...)")
answer_key_input = st.text_area("Paste Answer Key", height=150, help="Paste a single column of answers (A, B, C, D, etc.) from Excel or CSV. One per line.")

st.subheader("Answer Key Method")
answer_method = st.radio(
    "How would you like to handle answer keys?",
    ("Asterisk Method (mark correct answers with * in questions)", 
     "Answer Key Section (add separate answer list at end)"),
    help="ExamSoft supports both methods. Asterisk method marks correct answers directly in questions. Answer Key section adds a separate list at the end."
)
use_asterisk_method = answer_method.startswith("Asterisk")

# Initialize session state
if 'processed_data' not in st.session_state:
    st.session_state.processed_data = None

if st.button("Process Data"):
    if not questions_input.strip():
        st.error("Exam questions are required. Please paste your questions above.")
    else:
        # Parse answer key and clean encoding
        answer_key = [clean_text_encoding(a.strip()) for a in answer_key_input.strip().splitlines() if a.strip()]
        st.write(f"Answer key loaded: {len(answer_key)} answers")
        st.write(f"Using: {answer_method}")
        
        # Generate filenames based on course, section, and professor inputs
        instructions_filename = generate_filename(course_input, section_input, professor_input, "ins", "docx")
        exam_filename = generate_filename(course_input, section_input, professor_input, "exm", "rtf")
        
        # Use pasted instructions and questions, clean encoding
        instructions_text = clean_text_encoding(instructions_input.strip()) if instructions_input.strip() else ""
        questions_text = questions_input.strip()
        
        # Parse and format questions
        questions_list = parse_questions_from_text(questions_text, answer_key, use_asterisk_method)
        
        if questions_list:
            # Validation: count MC questions and compare to answer key
            mc_count = sum(1 for q in questions_list if not q.startswith("Type: E"))
            essay_count = sum(1 for q in questions_list if q.startswith("Type: E"))
            if mc_count != len(answer_key):
                st.warning(f"Warning: Found {mc_count} multiple choice questions but {len(answer_key)} answers in the answer key. Please check for missing or extra answers.")
            
            # Create downloadable files
            exam_rtf_content = None
            exam_rtf_bytes = None
            instructions_docx = None
            
            # Generate instructions if provided
            if instructions_text:
                instructions_docx = generate_instructions_docx(instructions_text)
            
            # Generate exam RTF
            exam_rtf_content = create_rtf_content(
                questions_list, 
                answer_key if not use_asterisk_method else None, 
                not use_asterisk_method
            )
            
            # Try LibreOffice Docker/Azure API conversion for best quality
            with tempfile.TemporaryDirectory() as tmpdir:
                docx_path = os.path.join(tmpdir, "ExamSoft_Export.docx")
                rtf_path = os.path.join(tmpdir, "ExamSoft_Export.rtf")
                generate_docx_with_questions(questions_list, '', docx_path)
                
                try:
                    # Use Docker/Flask API for conversion (Azure or local)
                    api_endpoint = get_converter_endpoint()
                    convert_docx_to_rtf_via_api(docx_path, rtf_path, api_url=api_endpoint)
                    with open(rtf_path, "rb") as f:
                        exam_rtf_bytes = f.read()
                    
                    # Show success message based on endpoint type
                    if is_using_azure() if AZURE_CONFIG_AVAILABLE else False:
                        st.success("✅ RTF generated using Azure LibreOffice API with clean encoding and best formatting fidelity.")
                    else:
                        st.success("✅ RTF generated using local LibreOffice Docker API with clean encoding and best formatting fidelity.")
                except Exception as e:
                    endpoint_type = "Azure" if (is_using_azure() if AZURE_CONFIG_AVAILABLE else False) else "local Docker"
                    st.error(f"❌ LibreOffice {endpoint_type} API conversion failed: {e}")
                    if not (is_using_azure() if AZURE_CONFIG_AVAILABLE else False):
                        st.info("💡 Make sure the Docker container is running on localhost:8080")
                        st.info("🌤️ Or deploy to Azure using the deployment script for better reliability")
                    st.info("🔄 Using basic RTF conversion as fallback.")
            
            # Store processed data in session state
            st.session_state.processed_data = {
                'instructions_text': instructions_text,
                'instructions_docx': instructions_docx,
                'instructions_filename': instructions_filename,
                'questions_list': questions_list,
                'exam_rtf_content': exam_rtf_content,
                'exam_rtf_bytes': exam_rtf_bytes,
                'exam_filename': exam_filename,
                'mc_count': mc_count,
                'essay_count': essay_count,
                'answer_key': answer_key,
                'use_asterisk_method': use_asterisk_method
            }
            
            st.success(f"Processed {len(questions_list)} questions successfully!")
            st.info(f"Found {mc_count} multiple choice questions and {essay_count} essay questions")
            
            # Show answer key info
            if use_asterisk_method:
                asterisk_count = sum(1 for q in questions_list for line in q.split('\n') if line.startswith('*'))
                st.info(f"Marked {asterisk_count} correct answers with asterisks")
            else:
                st.info(f"Answer key section with {len(answer_key)} answers will be added at the end")
        else:
            st.error("No questions were found or formatted")
            st.write("Debug - Raw exam questions preview:")
            st.text(questions_text[:1000])

# Display download options if data has been processed
if st.session_state.processed_data:
    data = st.session_state.processed_data
    
    st.markdown("---")
    st.subheader("Download Files")
    
    # Display preview
    if data['instructions_text']:
        st.subheader("Instructions Preview")
        st.text(data['instructions_text'][:500] + "..." if len(data['instructions_text']) > 500 else data['instructions_text'])
    
    st.subheader("Questions Preview (First 3)")
    for i, q in enumerate(data['questions_list'][:3]):
        st.text(q[:200] + "..." if len(q) > 200 else q)
        st.markdown("---")
    
    # Download checkboxes and buttons
    st.subheader("Select Files to Download")
    
    col1, col2 = st.columns(2)
    
    with col1:
        if data['instructions_docx']:
            download_instructions = st.checkbox("📄 Download Instructions (DOCX)", value=True)
            if download_instructions:
                st.download_button(
                    label="📄 Download Instructions",
                    data=data['instructions_docx'],
                    file_name=data['instructions_filename'],
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    use_container_width=True
                )
        else:
            st.info("No instructions to download (none were provided)")
    
    with col2:
        download_exam = st.checkbox("📝 Download Exam (RTF)", value=True)
        if download_exam:
            # Use Docker API version if available, otherwise fallback to basic
            if data['exam_rtf_bytes']:
                st.download_button(
                    label="📝 Download Exam (LibreOffice API)",
                    data=data['exam_rtf_bytes'],
                    file_name=data['exam_filename'],
                    mime="text/rtf",
                    use_container_width=True
                )
            else:
                st.download_button(
                    label="📝 Download Exam (Basic RTF)",
                    data=data['exam_rtf_content'],
                    file_name=data['exam_filename'],
                    mime="text/rtf",
                    use_container_width=True
                )
    
    # SharePoint Upload Section  
    if SHAREPOINT_INTEGRATION_AVAILABLE:
        st.markdown("---")
        st.subheader("📤 Upload to SharePoint (Optional)")
        
        # Check if user is authenticated (persistent auth handles token validation)
        is_authenticated = (
            st.session_state.get('sp_access_token') or 
            st.session_state.get('access_token') or 
            st.session_state.get('sharepoint_access_token')
        )
        
        if is_authenticated:
            # Show upload options
            st.write("✅ **Ready to upload!** Select files to upload to SharePoint:")
            
            upload_instructions_sp = st.checkbox("📄 Upload Instructions to SharePoint") if data['instructions_docx'] else False
            upload_exam_sp = st.checkbox("📝 Upload Exam to SharePoint", value=True)
            
            if st.button("🚀 Upload to SharePoint", use_container_width=True):
                try:
                    with st.spinner("Uploading to SharePoint..."):
                        access_token = (
                            st.session_state.get('sp_access_token') or 
                            st.session_state.get('access_token') or 
                            st.session_state.get('sharepoint_access_token')
                        )
                        
                        upload_results = []
                        
                        # Upload instructions if selected
                        if upload_instructions_sp and data['instructions_docx']:
                            success, result = upload_to_sharepoint_corrected(
                                access_token, 
                                data['instructions_docx'], 
                                data['instructions_filename']
                            )
                            upload_results.append(("Instructions", success, result))
                        
                        # Upload exam if selected
                        if upload_exam_sp:
                            # Debug: Show what keys are available
                            st.write("🔍 **Debug - Available data keys:**", list(data.keys()))
                            
                            # Use the correct RTF content key
                            rtf_content = data.get('exam_rtf_bytes') or data.get('exam_rtf_content')
                            
                            # Debug: Show what we found
                            st.write(f"🔍 **Debug - RTF content type:** {type(rtf_content)}")
                            st.write(f"🔍 **Debug - RTF content size:** {len(rtf_content) if rtf_content and hasattr(rtf_content, '__len__') else 'N/A'}")
                            
                            if rtf_content:
                                success, result = upload_to_sharepoint_corrected(
                                    access_token, 
                                    rtf_content, 
                                    data['exam_filename']
                                )
                                upload_results.append(("Exam", success, result))
                            else:
                                upload_results.append(("Exam", False, "No RTF content available"))
                        
                        # Show results
                        for file_type, success, result in upload_results:
                            if success:
                                st.success(f"✅ {file_type} uploaded successfully!")
                                if isinstance(result, dict) and 'url' in result:
                                    st.write(f"📁 **Uploaded to:** Exam Procedures / ExamSoft / Import")
                                    st.write(f"🔗 **File URL:** {result['url']}")
                                    if 'size' in result:
                                        st.write(f"📊 **Size:** {result['size']} bytes")
                            else:
                                st.error(f"❌ Failed to upload {file_type}: {result}")
                        
                        if all(success for _, success, _ in upload_results):
                            st.balloons()
                            st.success("🎉 All files uploaded successfully to **Exam Procedures / ExamSoft / Import**!")
                        
                except Exception as e:
                    st.error(f"Upload error: {str(e)}")
        else:
            st.info("🔐 **Sign in with Microsoft 365** in the sidebar to upload to SharePoint")
            st.write("Your authentication will persist for up to 90 days!")
    else:
        st.markdown("---")
        st.subheader("📤 Upload to SharePoint (Optional)")
        st.warning("⚠️ SharePoint functionality not available.")
        st.write("Please ensure the `sharepoint_integration.py` module is available.")
        with st.expander("SharePoint Upload Settings", expanded=False):
            st.write("🔐 **Secure Microsoft 365 Authentication**")
            st.write("Sign in with your Charleston School of Law Microsoft 365 account to upload files securely.")
            
            # Note about app registration
            st.info("📋 **IT Setup Required**: This feature requires app registration in your Microsoft 365 tenant. Contact IT to configure the client ID and permissions.")
            
            # For now, show configuration fields that IT would need to set up
            st.write("**Required IT Configuration:**")
            st.code(f"""
App Registration Settings:
- Client ID: {M365_CONFIG['client_id']}
- Tenant: {M365_CONFIG['tenant_id']}
- Redirect URI: {M365_CONFIG['redirect_uri']}
- Required Permissions:
  • Sites.ReadWrite.All
  • Files.ReadWrite.All
            """)
            
            # Placeholder authentication button (would work once app is registered)
            if st.button("� Sign in with Microsoft 365"):
                st.warning("⚠️ **App Registration Required**: Contact your IT department to register this application with Microsoft 365 before using SharePoint integration.")
                st.write("Once configured, this button will:")
                st.write("• Open Microsoft 365 login in your browser")
                st.write("• Securely authenticate without passwords")
                st.write("• Show available SharePoint sites and libraries")
                st.write("• Allow direct file uploads")
            
            # Mock interface to show what it would look like when working
            st.write("---")
            st.write("**Preview of Full SharePoint Integration** (once IT setup is complete):")
            
            # Upload checkboxes (currently disabled)
            upload_instructions_sp = st.checkbox("📄 Upload Instructions to SharePoint", disabled=True) if data['instructions_docx'] else False
            upload_exam_sp = st.checkbox("📝 Upload Exam to SharePoint", value=True, disabled=True)
            
            st.button("🚀 Upload to SharePoint", disabled=True, help="Requires Microsoft 365 app registration")
    
    # Action buttons after processing
    st.markdown("---")
    col1, col2, col3 = st.columns([1, 2, 1])
    
    with col1:
        if st.button("🔄 Convert Another Exam", use_container_width=True, help="Clear form and start over"):
            # Clear all processed data and form inputs
            st.session_state.processed_data = None
            # Clear form session state if it exists
            for key in list(st.session_state.keys()):
                if key.startswith('form_') or key in ['course_input', 'section_input', 'professor_input']:
                    del st.session_state[key]
            st.success("📝 Form cleared! Ready for next exam.")
            # Don't use st.rerun() - let natural refresh handle it
    
    with col3:
        st.write("") # Spacer

else:
    st.info("Paste your exam questions and answer key above, then click 'Process Data'. Instructions are optional.")


# This file contains the main Streamlit UI code
# The UI runs immediately when imported
