"""
Minimal RTF formatter - only the core functions needed for the app
"""
import re
import tempfile
import os
from io import BytesIO
from datetime import datetime
from docx import Document


def clean_text_encoding(text):
    """Clean and normalize text encoding to handle copy/paste issues"""
    if not text:
        return text
    
    # Common encoding fixes
    replacements = {
        ''': "'", ''': "'", '"': '"', '"': '"', '–': '-', '—': '--',
        '…': '...', '•': '- ', '®': '(R)', '©': '(C)', '™': '(TM)',
        '½': '1/2', '¼': '1/4', '¾': '3/4', '°': ' degrees', 'é': 'e',
        'è': 'e', 'ê': 'e', 'ë': 'e', 'à': 'a', 'á': 'a', 'â': 'a',
        'ã': 'a', 'ä': 'a', 'å': 'a', 'ç': 'c', 'ì': 'i', 'í': 'i',
        'î': 'i', 'ï': 'i', 'ñ': 'n', 'ò': 'o', 'ó': 'o', 'ô': 'o',
        'õ': 'o', 'ö': 'o', 'ø': 'o', 'ù': 'u', 'ú': 'u', 'û': 'u',
        'ü': 'u', 'ý': 'y', 'ÿ': 'y'
    }
    
    for old, new in replacements.items():
        text = text.replace(old, new)
    
    # Clean up extra whitespace
    text = re.sub(r'\s+', ' ', text)
    text = text.strip()
    
    return text


def parse_questions_from_text(text, answer_key=None, use_asterisk_method=False):
    """Parse questions from text input"""
    if not text.strip():
        return []
    
    questions = []
    lines = text.split('\n')
    current_question = []
    question_counter = 0
    
    for line in lines:
        line = line.strip()
        if not line:
            continue
            
        # Check if this is a new question (starts with number)
        if re.match(r'^\d+[\.\)]\s*', line):
            # Save previous question
            if current_question:
                question_text = '\n'.join(current_question)
                
                # Apply asterisk method if enabled
                if use_asterisk_method and answer_key and question_counter < len(answer_key):
                    correct_answer = answer_key[question_counter].upper()
                    question_text = apply_asterisk_to_question(question_text, correct_answer)
                
                questions.append(question_text)
                question_counter += 1
            
            # Start new question
            current_question = [line]
        else:
            # Continue current question
            if current_question:
                current_question.append(line)
    
    # Add last question
    if current_question:
        question_text = '\n'.join(current_question)
        if use_asterisk_method and answer_key and question_counter < len(answer_key):
            correct_answer = answer_key[question_counter].upper()
            question_text = apply_asterisk_to_question(question_text, correct_answer)
        questions.append(question_text)
    
    return questions


def apply_asterisk_to_question(question_text, correct_answer):
    """Apply asterisk to the correct answer choice"""
    lines = question_text.split('\n')
    processed_lines = []
    
    for line in lines:
        # Check if this line is an answer choice
        if re.match(r'^[A-Z][\.\)]\s*', line):
            choice_letter = line[0].upper()
            if choice_letter == correct_answer:
                # Add asterisk to correct answer
                line = '*' + line
        processed_lines.append(line)
    
    return '\n'.join(processed_lines)


def generate_filename(course, section, professor, file_type, extension):
    """Generate standardized filename"""
    parts = []
    if course:
        parts.append(course.upper())
    if section:
        parts.append(section)
    if professor:
        parts.append(professor.title())
    
    if parts:
        base_name = '_'.join(parts) + f'_{file_type}'
    else:
        base_name = f'ExamSoft_{file_type}_{datetime.now().strftime("%Y%m%d")}'
    
    return f"{base_name}.{extension}"


def create_rtf_content(questions_list, answer_key=None, include_answer_section=False):
    """Create RTF content from questions"""
    if not questions_list:
        return ""
    
    rtf_header = r"{\rtf1\ansi\deff0 {\fonttbl {\f0 Times New Roman;}} \f0\fs24 "
    rtf_body = ""
    
    # Add questions
    for i, question in enumerate(questions_list, 1):
        rtf_body += f"\\par\\par {question}\\par "
    
    # Add answer key section if requested
    if include_answer_section and answer_key:
        rtf_body += "\\par\\par\\b ANSWER KEY:\\b0\\par "
        for i, answer in enumerate(answer_key, 1):
            rtf_body += f"{i}. {answer.upper()}\\par "
    
    return rtf_header + rtf_body + "}"


def generate_instructions_docx(instructions_text):
    """Generate instructions DOCX file"""
    if not instructions_text:
        return None
    
    doc = Document()
    doc.add_heading('Exam Instructions', 0)
    doc.add_paragraph(clean_text_encoding(instructions_text))
    
    # Save to bytes
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()


def generate_docx_with_questions(questions_list, instructions_text, output_path):
    """Generate DOCX file with questions"""
    doc = Document()
    
    if instructions_text:
        doc.add_heading('Instructions', 1)
        doc.add_paragraph(clean_text_encoding(instructions_text))
        doc.add_page_break()
    
    doc.add_heading('Exam Questions', 1)
    
    for question in questions_list:
        doc.add_paragraph(clean_text_encoding(question))
        doc.add_paragraph()  # Add space between questions
    
    doc.save(output_path)


def get_converter_endpoint():
    """Get converter endpoint URL"""
    return "http://localhost:8080/convert"  # Default local endpoint


def is_using_azure():
    """Check if using Azure endpoint"""
    return False  # Default to local


def convert_docx_to_rtf_via_api(docx_path, rtf_path, api_url="http://localhost:8080/convert"):
    """Convert DOCX to RTF using API"""
    import requests
    
    with open(docx_path, 'rb') as f:
        files = {'file': f}
        response = requests.post(api_url, files=files, timeout=30)
        response.raise_for_status()
        
        with open(rtf_path, 'wb') as output_file:
            output_file.write(response.content)


def upload_to_sharepoint_corrected(access_token, file_content, filename):
    """Upload file to SharePoint"""
    # This is a placeholder - would need actual SharePoint implementation
    return False, "SharePoint upload not implemented in minimal version"